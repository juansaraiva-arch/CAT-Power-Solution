"""
CAT Power Solution — Proposal Document Generator
=================================================
Generates a customer proposal Word document (.docx) by combining sizing results,
project header info, and commercial/administrative fields.

This module does NOT modify any sizing logic — it is purely a document builder.
"""

import io
import os
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from core.generator_library import GENERATOR_LIBRARY


# ---------------------------------------------------------------------------
# Styling constants
# ---------------------------------------------------------------------------
_FONT_NAME = "Arial"
_CAT_YELLOW = "FFCC00"
_TABLE_HEADER_BG = "D3D1C7"
_TABLE_BORDER_COLOR = "AAAAAA"
_HEADER_GRAY = RGBColor(0x99, 0x99, 0x99)

_PAGE_WIDTH = Inches(8.5)
_PAGE_HEIGHT = Inches(11)
_MARGIN = Inches(1)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _safe(obj, attr, fallback="TBD"):
    """Safely get an attribute from an object, returning *fallback* on failure."""
    try:
        val = getattr(obj, attr)
        if val is None:
            return fallback
        return val
    except (AttributeError, TypeError):
        return fallback


def _safe_dict(d, key, fallback="TBD"):
    """Safely get a key from a dict."""
    if not isinstance(d, dict):
        return fallback
    val = d.get(key)
    return val if val is not None else fallback


def _set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_paragraph_shading(paragraph, color_hex):
    """Apply background shading to a paragraph."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{color_hex}"/>')
    paragraph._p.get_or_add_pPr().append(shading)


def _set_run_font(run, size_pt, bold=False, color=None, name=_FONT_NAME):
    """Configure a run's font properties."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_paragraph(doc, text, size=11, bold=False, alignment=None, space_before=None, space_after=None):
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size, bold=bold)
    if alignment is not None:
        p.alignment = alignment
    fmt = p.paragraph_format
    fmt.line_spacing = 1.15
    if space_before is not None:
        fmt.space_before = Pt(space_before)
    if space_after is not None:
        fmt.space_after = Pt(space_after)
    return p


def _add_heading(doc, text, level=1):
    """Add a heading with correct font styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = _FONT_NAME
        if level == 1:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        run.bold = True
    fmt = h.paragraph_format
    if level == 1:
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(6)
    else:
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(3)
    return h


def _set_table_borders(table):
    """Apply thin borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{_TABLE_BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{_TABLE_BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{_TABLE_BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{_TABLE_BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{_TABLE_BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{_TABLE_BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _style_header_row(row):
    """Bold text, gray background, 10pt for a table header row."""
    for cell in row.cells:
        _set_cell_shading(cell, _TABLE_HEADER_BG)
        for p in cell.paragraphs:
            for run in p.runs:
                _set_run_font(run, 10, bold=True)


def _style_body_cell(cell):
    """10pt normal for table body cells."""
    for p in cell.paragraphs:
        for run in p.runs:
            _set_run_font(run, 10)


def _add_bullet(doc, text, size=11):
    """Add a bullet-point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    # Clear default run and add styled one
    p.clear()
    run = p.add_run(text)
    _set_run_font(run, size)
    return p


def _offer_checkboxes(proposal_info):
    """Return offer-type checkbox string with 2-column layout."""
    genset_only = proposal_info.get("offer_type_genset", False)
    switchgear = proposal_info.get("offer_type_switchgear", False)
    energy_storage = proposal_info.get("offer_type_energy_storage", False)
    solution_enclosure = proposal_info.get("offer_type_solutions", False)
    scr = proposal_info.get("offer_type_scr", False)
    other = proposal_info.get("offer_type_other", proposal_info.get("offer_type_hybrid", False))

    def cb(checked):
        return "\u2611" if checked else "\u2610"

    col1 = f"[{cb(genset_only)} Genset Only]"
    col2_items = [
        f"[{cb(switchgear)} Switchgear]",
        f"[{cb(energy_storage)} Energy Storage]",
        f"[{cb(solution_enclosure)} Solution/Enclosure]",
        f"[{cb(scr)} Selective Catalytic Reduction Solution (SCR)]",
        f"[{cb(other)} Other]",
    ]
    col2 = "  ".join(col2_items)
    return f"{col1}\t\tBalance of Plant (BOP) Elements: {col2}"


_LOGO_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "assets", "logo_caterpillar.png")


def _add_page_header(section, project_name):
    """Add running header with logo image (left) and page number (right)."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    # Left-aligned: logo image + project name
    run_logo = p.add_run()
    if os.path.exists(_LOGO_PATH):
        run_logo.add_picture(_LOGO_PATH, height=Inches(0.3))
    run_sep = p.add_run(f"  \u2014  {project_name}")
    _set_run_font(run_sep, 9, color=_HEADER_GRAY)
    # Tab to push page number right
    p.add_run("\t\t")
    # Page number field
    run_pg = p.add_run()
    _set_run_font(run_pg, 9, color=_HEADER_GRAY)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_pg._r.append(fldChar1)
    run_pg2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_pg2._r.append(instrText)
    run_pg3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_pg3._r.append(fldChar2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ---------------------------------------------------------------------------
# Main generator function
# ---------------------------------------------------------------------------

def _generate_proposal_docx_legacy(
    sizing_result,
    header_info: dict,
    proposal_info: dict,
    output_path: str = None,
) -> bytes:
    """
    Legacy proposal generator (used by render_docx_proposal_tab / ENABLE_PROPOSAL_GEN).

    Parameters
    ----------
    sizing_result : object
        SizingResult from the sizing engine. All attribute accesses are guarded.
    header_info : dict
        project_name, client_name, contact_name, contact_email, contact_phone,
        country, state_province.
    proposal_info : dict
        Fields from PROPOSAL_DEFAULTS with user overrides applied.
    output_path : str, optional
        If given, writes the file to disk AND returns bytes.

    Returns
    -------
    bytes
        The .docx file content.
    """
    doc = Document()

    # --- Page setup -----------------------------------------------------------
    section = doc.sections[0]
    section.page_width = _PAGE_WIDTH
    section.page_height = _PAGE_HEIGHT
    section.top_margin = _MARGIN
    section.bottom_margin = _MARGIN
    section.left_margin = _MARGIN
    section.right_margin = _MARGIN

    project_name = header_info.get("project_name", "Untitled Project")
    client_name = header_info.get("client_name", "")
    dealer_name = proposal_info.get("dealer_name", "")

    # --- Lookup generator specs -----------------------------------------------
    gen_model = _safe(sizing_result, "selected_gen", "TBD")
    gen_specs = GENERATOR_LIBRARY.get(gen_model, {})
    iso_rating = gen_specs.get("iso_rating_mw", "TBD")
    freq_hz = _safe(sizing_result, "freq_hz", gen_specs.get("voltage_kv", "TBD"))
    voltage_kv = gen_specs.get("voltage_kv", "TBD")

    # Try to get freq_hz from sizing_result first
    freq_hz = _safe(sizing_result, "freq_hz", "TBD")

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    # CATERPILLAR logo (cover page)
    p_cat = doc.add_paragraph()
    p_cat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo_cover = p_cat.add_run()
    if os.path.exists(_LOGO_PATH):
        run_logo_cover.add_picture(_LOGO_PATH, width=Inches(3.5))
    else:
        # Fallback to text if logo file missing
        run_logo_cover.text = "CATERPILLAR"
        _set_run_font(run_logo_cover, 16, bold=True)

    # Division
    _add_paragraph(doc, proposal_info.get("cat_division", ""), 12,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Proposal type
    _add_paragraph(doc, proposal_info.get("proposal_type", ""), 12,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Customer Proposal For:
    _add_paragraph(doc, "Customer Proposal For:", 12, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Project name
    _add_paragraph(doc, project_name, 14, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Generator model
    _add_paragraph(doc, str(gen_model), 12,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Date
    today_str = date.today().strftime("%B %d, %Y")
    _add_paragraph(doc, today_str, 11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Version
    _add_paragraph(doc, proposal_info.get("doc_version", "Rev. 0"), 11,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # --- Add running header from second page onward ---------------------------
    new_section = doc.add_section()
    new_section.page_width = _PAGE_WIDTH
    new_section.page_height = _PAGE_HEIGHT
    new_section.top_margin = _MARGIN
    new_section.bottom_margin = _MARGIN
    new_section.left_margin = _MARGIN
    new_section.right_margin = _MARGIN
    _add_page_header(new_section, project_name)

    # =========================================================================
    # SECTION 1 — EXECUTIVE SUMMARY
    # =========================================================================
    _add_heading(doc, "Executive Summary", level=1)

    exec_summary = (
        f"Caterpillar Inc. is pleased to present this proposal for the {project_name}. "
        "The solution offered in this document is designed to meet the project\u2019s operational, "
        "performance, and reliability requirements while providing a flexible and scalable "
        "foundation for long-term use. The proposed equipment and supporting solutions are based "
        "on proven designs widely deployed across diverse Applications and operating environments."
    )
    _add_paragraph(doc, exec_summary)

    exec_summary_2 = (
        "This document outlines a preliminary, non-binding proposal based on the information "
        "provided to date. Final pricing and commercial terms will be confirmed and executed "
        "through the authorized Caterpillar dealer."
    )
    _add_paragraph(doc, exec_summary_2)

    exec_summary_3 = (
        "Our organization has extensive experience supplying electric power solutions for "
        "industrial, commercial, and critical-infrastructure projects around the world. This "
        "background ensures that the solution outlined in this proposal reflects globally "
        "recognized best practices in safety, performance, and integration with customer "
        "facilities. The equipment included in this offer is supported by comprehensive "
        "technical capabilities, a mature product line, and global service resources to assist "
        "throughout the lifecycle of the installation."
    )
    _add_paragraph(doc, exec_summary_3)

    exec_summary_4 = (
        "This proposal includes the supply of the primary equipment and associated controls, "
        "along with supporting solutions required for reliable operation. A detailed description "
        "of the technical configuration, performance parameters, and solution interfaces is "
        "provided in the following sections. This offer has been configured to align with the "
        "project information provided and can be adapted to meet additional requirements as needed."
    )
    _add_paragraph(doc, exec_summary_4)

    exec_summary_5 = (
        f"Commercial information, including pricing, delivery expectations, validity, and warranty "
        f"provisions, is included within this document. All pricing and commercial conditions are "
        f"presented based on the defined scope and assumptions associated with the project. The "
        f"proposal is intended to provide {client_name} with a clear understanding of the recommended "
        f"solution and the associated commercial framework."
    )
    _add_paragraph(doc, exec_summary_5)

    exec_summary_6 = (
        "We appreciate the opportunity to be considered for this project and trust that the "
        "information presented will support your evaluation of our proposed electric power solution. "
        "Our team is available to discuss the contents of this proposal in further detail and to "
        "assist in identifying the configuration best suited for your Application."
    )
    _add_paragraph(doc, exec_summary_6)

    # Closing
    bdm_name = proposal_info.get("bdm_name", "")
    bdm_email = proposal_info.get("bdm_email", "")
    cat_division = proposal_info.get("cat_division", "")

    _add_paragraph(doc, "Sincerely,", 11, space_before=12)
    _add_paragraph(doc, bdm_name, 11)
    _add_paragraph(doc, bdm_email, 11)
    _add_paragraph(doc, f"Caterpillar Inc. \u2014 {cat_division}", 11)

    doc.add_page_break()

    # =========================================================================
    # SECTION 2 — SOLUTION OVERVIEW
    # =========================================================================
    _add_heading(doc, "Proposed Customer Solution", level=1)

    # Total facility MW
    total_facility_mw = _safe(sizing_result, "p_total_dc",
                              _safe(sizing_result, "total_facility_mw", "TBD"))
    try:
        total_facility_mw_str = f"{float(total_facility_mw):.1f}"
    except (ValueError, TypeError):
        total_facility_mw_str = str(total_facility_mw)

    # Application type
    try:
        application = sizing_result.dc_type
        if not application:
            application = "Data Center"
    except AttributeError:
        application = "Data Center"

    narrative_s2 = (
        f"Caterpillar Inc. is pleased to provide this proposed power generation solution for "
        f"{client_name}, tailored to support the operational requirements of the {project_name}."
    )
    _add_paragraph(doc, narrative_s2)

    narrative_s2b = (
        f"This proposal aims to satisfy the request for a {total_facility_mw_str} MW, "
        f"{application} generation site utilizing the {gen_model} Caterpillar Genset."
    )
    _add_paragraph(doc, narrative_s2b)

    # Bullet list
    # N units
    n_units = _safe(sizing_result, "n_total", "TBD")

    # Derated MW per unit
    derated_mw = _safe(sizing_result, "unit_site_cap", "TBD")
    try:
        derated_mw_str = f"{float(derated_mw):.2f}"
    except (ValueError, TypeError):
        derated_mw_str = str(derated_mw)

    _add_bullet(doc, f"{n_units} x {gen_model} Caterpillar generator set solutions")
    _add_bullet(doc, f"Derated output per unit: {derated_mw_str} MW at site conditions")

    # BESS
    use_bess = _safe(sizing_result, "use_bess", False)
    bess_power = _safe(sizing_result, "bess_power_mw", 0)
    bess_energy = _safe(sizing_result, "bess_energy_mwh", 0)
    if use_bess and bess_power:
        try:
            bess_str = f"Yes \u2014 {float(bess_power):.1f} MW / {float(bess_energy):.1f} MWh"
        except (ValueError, TypeError):
            bess_str = "Not included"
    else:
        bess_str = "Not included"
    _add_bullet(doc, f"BESS: {bess_str}")

    # Availability
    avail = _safe(sizing_result, "system_availability", "TBD")
    try:
        avail_str = f"{float(avail) * 100:.2f}%"
    except (ValueError, TypeError):
        avail_str = str(avail)

    n_running = _safe(sizing_result, "n_running", "")
    n_reserve = _safe(sizing_result, "n_reserve", "")
    config_label = f"N+{n_reserve}" if n_reserve != "" else ""
    _add_bullet(doc, f"System availability: {avail_str} ({config_label} configuration)")
    _add_bullet(doc, "CVA and ESC options available \u2014 refer to Exhibit for overview")

    doc.add_page_break()

    # =========================================================================
    # SECTION 3 — TECHNICAL OFFER
    # =========================================================================
    _add_heading(doc, "Technical Offer", level=1)

    # --- 3.1 Equipment Summary ------------------------------------------------
    _add_heading(doc, "3.1 Equipment Summary", level=2)
    _add_paragraph(doc, (
        "The Equipment Summary below outlines the equipment and selected features "
        "included in the proposed power generation solution."
    ))

    # Offer type checkboxes
    _add_paragraph(doc, _offer_checkboxes(proposal_info), 11)

    # Equipment table
    eq_table = doc.add_table(rows=1, cols=4)
    eq_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = eq_table.rows[0]
    for i, label in enumerate(["Qty.", "Characteristic Name", "Feature Code", "Feature Description"]):
        hdr.cells[i].text = label
    _style_header_row(hdr)
    _set_table_borders(eq_table)

    # Row 1 — generator
    try:
        feat_desc = f"{iso_rating} MW ISO, {freq_hz} Hz, {voltage_kv} kV"
    except Exception:
        feat_desc = "TBD"
    row1 = eq_table.add_row()
    row1.cells[0].text = str(n_units)
    row1.cells[1].text = "Generator Set"
    row1.cells[2].text = str(gen_model)
    row1.cells[3].text = feat_desc
    for c in row1.cells:
        _style_body_cell(c)

    # Row 2 — BESS (conditional)
    if use_bess and bess_power:
        row_bess = eq_table.add_row()
        row_bess.cells[0].text = "1"
        row_bess.cells[1].text = "Battery Energy Storage System"
        row_bess.cells[2].text = "BESS"
        try:
            row_bess.cells[3].text = f"{float(bess_power):.1f} MW / {float(bess_energy):.1f} MWh"
        except (ValueError, TypeError):
            row_bess.cells[3].text = "TBD"
        for c in row_bess.cells:
            _style_body_cell(c)

    # 3 empty rows for BDM
    for _ in range(3):
        empty_row = eq_table.add_row()
        for c in empty_row.cells:
            c.text = ""
            _style_body_cell(c)

    # --- 3.2 Project Requirements ---------------------------------------------
    _add_heading(doc, "3.2 Project Requirements", level=2)
    _add_paragraph(doc, "The following project parameters were considered in defining the technical offering:")

    country = header_info.get("country", "To be confirmed")
    state_province = header_info.get("state_province", "To be confirmed")
    _add_bullet(doc, f"Project location / site: {country}, {state_province}")

    _add_bullet(doc, "Intended application: \u2611 Prime  \u2610 Standby  \u2610 Other")

    # Site conditions
    site_temp = _safe(sizing_result, "site_temp_c",
                      _safe_dict(_safe(sizing_result, "sizing_input", {}),
                                 "site_temp_c", "To be confirmed"))
    site_alt = _safe(sizing_result, "site_alt_m",
                     _safe_dict(_safe(sizing_result, "sizing_input", {}),
                                "site_alt_m", "To be confirmed"))
    _add_bullet(doc, f"Operating environment: Ambient temperature {site_temp}\u00b0C, Altitude {site_alt} m ASL")

    # Load requirements
    p_it = _safe(sizing_result, "p_it", "To be confirmed")
    pue = _safe(sizing_result, "pue", "To be confirmed")
    step_load = _safe(sizing_result, "step_load_pct",
                      gen_specs.get("step_load_pct", "To be confirmed"))
    _add_bullet(doc, (
        f"Load requirements: IT Load {p_it} MW, PUE {pue}, "
        f"Total Facility {total_facility_mw_str} MW, Step load {step_load}%"
    ))

    acoustic = _safe(sizing_result, "acoustic_treatment",
                     _safe_dict(_safe(sizing_result, "noise_results", {}),
                                "acoustic_treatment", "To be confirmed"))
    _add_bullet(doc, f"Installation conditions: {acoustic} acoustic treatment")

    # --- 3.3 Fuel Requirements ------------------------------------------------
    _add_heading(doc, "3.3 Fuel Requirements", level=2)

    fuel_mode = _safe(sizing_result, "fuel_mode",
                      _safe_dict(_safe(sizing_result, "sizing_input", {}),
                                 "fuel_mode", "To be confirmed"))
    methane_number = _safe(sizing_result, "methane_number",
                           _safe_dict(_safe(sizing_result, "sizing_input", {}),
                                      "methane_number", "To be confirmed"))

    _add_paragraph(doc, f"The proposed generator set is designed to operate using {fuel_mode}.")
    _add_bullet(doc, f"Fuel type: {fuel_mode}")
    _add_bullet(doc, f"Methane number: {methane_number}")
    _add_bullet(doc, "Estimated consumption at full load: To be confirmed per fuel analysis")
    _add_bullet(doc, "Note: Detailed fuel specifications are included in Exhibit H.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 4 — PRICING
    # =========================================================================
    _add_heading(doc, "Pricing", level=1)

    # --- 4.1 Base Price -------------------------------------------------------
    _add_heading(doc, "4.1 Base Price", level=2)

    capex_val = _safe(sizing_result, "total_capex", None)
    if capex_val is None:
        # Try financial.capex_total_musd pattern
        try:
            capex_val = sizing_result.financial.capex_total_musd * 1_000_000
        except (AttributeError, TypeError):
            capex_val = None

    if capex_val is not None:
        try:
            formatted_capex = f"${float(capex_val):,.0f}"
        except (ValueError, TypeError):
            formatted_capex = "To be confirmed"
    else:
        formatted_capex = "To be confirmed"

    _add_paragraph(doc, f"Base Price: {formatted_capex}", 11, bold=True)
    _add_paragraph(doc, "Reflects the fully integrated system scope outlined in this proposal.")
    _add_paragraph(doc, (
        f"This is a budgetary estimate. Final pricing to be confirmed through "
        f"authorized CAT dealer: {dealer_name}."
    ))

    # --- 4.2 Additional Options -----------------------------------------------
    _add_heading(doc, "4.2 Additional Options", level=2)
    _add_paragraph(doc, (
        "The following optional features are available upon request and are "
        "not included in the Base Price."
    ))

    additional = proposal_info.get("additional_options", [])
    opt_table = doc.add_table(rows=1, cols=3)
    opt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_opt = opt_table.rows[0]
    for i, label in enumerate(["Option", "Description", "Price (USD)"]):
        hdr_opt.cells[i].text = label
    _style_header_row(hdr_opt)
    _set_table_borders(opt_table)

    for item in additional:
        r = opt_table.add_row()
        r.cells[0].text = str(item.get("description", ""))
        r.cells[1].text = str(item.get("description", ""))
        try:
            r.cells[2].text = f"${float(item.get('price_usd', 0)):,.0f}"
        except (ValueError, TypeError):
            r.cells[2].text = ""
        for c in r.cells:
            _style_body_cell(c)

    # Ensure at least 3 body rows
    rows_needed = max(0, 3 - len(additional))
    for _ in range(rows_needed):
        empty = opt_table.add_row()
        for c in empty.cells:
            c.text = ""
            _style_body_cell(c)

    # --- 4.3 Price Validity ---------------------------------------------------
    _add_heading(doc, "4.3 Price Validity", level=2)
    validity = proposal_info.get("proposal_validity", "90 days from date of issue")
    _add_paragraph(doc, f"This proposal is valid until {validity}.")

    # --- 4.4 Payment Terms ----------------------------------------------------
    _add_heading(doc, "4.4 Payment Terms", level=2)
    _add_paragraph(doc, "The proposed payment terms are as follows:")
    _add_bullet(doc, f"{proposal_info.get('payment_pct_down', 30)}% downpayment")
    _add_bullet(doc, f"{proposal_info.get('payment_pct_30d', 30)}% due 30 days after purchase order")
    _add_bullet(doc, f"{proposal_info.get('payment_pct_sol', 30)}% due at SOL (Start on Line)")
    _add_bullet(doc, f"{proposal_info.get('payment_pct_rts', 10)}% due at written notice of RTS (Ready to Ship)")

    # --- 4.5 Pricing Parameters -----------------------------------------------
    _add_heading(doc, "4.5 Pricing Parameters", level=2)
    _add_paragraph(doc, (
        "a) The indicated price is based on the quoted delivery. Pricing may change "
        "if the customer requests a delay in delivery."
    ))
    _add_paragraph(doc, (
        "b) We reserve the right to adjust our prices in case some parts of our offer "
        "are inapplicable and will not be ordered."
    ))
    _add_paragraph(doc, (
        f"The terms of this proposal are issued through {dealer_name}, and the final "
        f"contract will be executed between {dealer_name} and the Customer."
    ))

    # --- 4.6 Taxes, Licenses, and Fees ----------------------------------------
    _add_heading(doc, "4.6 Taxes, Licenses, and Fees", level=2)
    incoterm = proposal_info.get("incoterm", "Ex Works (EXW)")

    _add_paragraph(doc, (
        "a) All equipment duties, taxes of any kind, and license fees \u2014 including VAT "
        "and/or sales tax imposed by national or local government \u2014 are not included in "
        "this quotation and will be the responsibility of the Customer where applicable."
    ))
    _add_paragraph(doc, f"b) Delivery under {incoterm} per Incoterms.")
    _add_paragraph(doc, (
        "c) The Customer shall be responsible to secure all relevant construction permits "
        "from the national and local government agencies, and all the actual fees shall be "
        "paid by the Customer."
    ))

    # --- 4.7 Delivery ---------------------------------------------------------
    _add_heading(doc, "4.7 Delivery", level=2)
    delivery_date = proposal_info.get("delivery_date_est", "To be confirmed")
    delivery_dest = proposal_info.get("delivery_destination", "Ex Factory")

    _add_paragraph(doc, (
        f"a) With a Purchase Order, or document evidencing financial commitment (LOA) within "
        f"the validity period stated above, the units will be available for delivery by "
        f"approximately {delivery_date}."
    ))
    _add_paragraph(doc, f"b) Delivery destination: {delivery_dest}")

    doc.add_page_break()

    # =========================================================================
    # SECTION 5 — CLARIFICATIONS, DEVIATIONS & EXCEPTIONS
    # =========================================================================
    _add_heading(doc, "Clarifications, Deviations & Exceptions", level=1)

    _add_paragraph(doc, _offer_checkboxes(proposal_info), 11)

    clar_table = doc.add_table(rows=1, cols=1)
    clar_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    clar_table.rows[0].cells[0].text = "Project Clarifications, Deviations & Exceptions"
    _style_header_row(clar_table.rows[0])
    _set_table_borders(clar_table)

    boilerplate_row = clar_table.add_row()
    boilerplate_row.cells[0].text = (
        "Equipment supplied will be limited to that described in this proposal. If products "
        "and/or features, other than the equipment described in this proposal is required, "
        "please consult the factory for a re-quote as pricing may be affected. We take "
        "complete exception to any specifications that were not provided or reviewed."
    )
    _style_body_cell(boilerplate_row.cells[0])

    notes = proposal_info.get("proposal_notes", "")
    if notes:
        notes_row = clar_table.add_row()
        notes_row.cells[0].text = notes
        _style_body_cell(notes_row.cells[0])

    doc.add_page_break()

    # =========================================================================
    # SECTION 6 — COMMERCIAL ASSUMPTIONS
    # =========================================================================
    _add_heading(doc, "Commercial Assumptions", level=1)

    _add_paragraph(doc, _offer_checkboxes(proposal_info), 11)

    assumptions = [
        "All prices exclude all duties, taxes, trade surcharges, and freight.",
        (
            "Receipt of goods shall be confirmed by the customer. Notification and evidence "
            "of any material damaged during shipment and/or any concerns with the products "
            "shall be provided as soon as possible but not later than 10 business days "
            "following shipment loading or 5 days after receipt, whichever comes first."
        ),
        (
            "Unless specifically indicated, all job site installation, commissioning and "
            "testing for the project is NOT included."
        ),
        (
            "Unless specifically indicated, storage of offered materials at the manufacturing "
            "location is NOT available. Completed units must be picked up from the manufacturing "
            "facility within 5 business days of RTS. Units that are not picked up within 5 "
            "business days will be moved to an off-site storage facility at customer\u2019s expense."
        ),
        (
            "Any modification to the defined scope of supply requires a mutually agreed and "
            "signed change order."
        ),
        (
            "Caterpillar product warranty is governed by the applicable Caterpillar Inc. "
            "Warranty Guide(s)."
        ),
        (
            "Orders may be cancelled only by written notice. Payments made under Payment Terms "
            "are non-refundable and will not be returned in event of cancellation."
        ),
    ]

    assump_table = doc.add_table(rows=1, cols=2)
    assump_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_a = assump_table.rows[0]
    hdr_a.cells[0].text = "No."
    hdr_a.cells[1].text = "Description"
    _style_header_row(hdr_a)
    _set_table_borders(assump_table)

    for idx, text in enumerate(assumptions, 1):
        row = assump_table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = text.replace("[DEALER NAME]", dealer_name)
        for c in row.cells:
            _style_body_cell(c)

    doc.add_page_break()

    # =========================================================================
    # SECTION 7 — APPENDICES
    # =========================================================================

    # --- Appendix A — Definitions ---------------------------------------------
    _add_heading(doc, "Exhibit A \u2014 Definitions", level=1)

    definitions = [
        ("Application (Standby / Prime / Other)",
         "The operational mode of the generator set. Standby is for emergency use only during "
         "utility outages, while Prime is intended for continuous or frequent operation as the "
         "main power source."),
        ("Balance of Plant (BOP)",
         "All supporting equipment required for a complete power solution beyond the generator set "
         "itself. In this Proposal, BOP includes switchgear, enclosures/solutions, SCR solutions, "
         "inverters, energy storage, and other ancillary components."),
        ("Base Price",
         "The total price for the defined Scope of Supply included in the Proposal, excluding "
         "optional items, taxes, duties, freight, and licensing fees."),
        ("Bill of Material (BOM)",
         "A summarized list of all equipment, components, features, and selected options included "
         "in the proposed power generation solution."),
        ("Clarification",
         "An explanation of assumptions or interpretations where project information is incomplete "
         "or ambiguous, provided to ensure shared understanding of the proposal."),
        ("Delivery",
         "A Delivery condition in which Caterpillar makes the equipment available at the factory, "
         "and the Customer assumes responsibility for transportation, duties, permits, and logistics "
         "from that point forward."),
        ("Deviation",
         "A departure from the customer\u2019s specification when the proposed solution cannot fully "
         "conform to the requirements provided."),
        ("Exception",
         "A requested requirement or obligation that Caterpillar cannot comply with or accept as "
         "part of this proposal."),
        ("Feature Code",
         "The identifier used to specify a configurable feature, option, or characteristic of the "
         "generator set or BOP equipment."),
        ("Fuel Requirements",
         "The fuel type and characteristics required for operation of the generator set, such as "
         "natural gas, biogas, diesel, or propane."),
        ("Load Requirements",
         "The expected electrical demand characteristics of the project, including maximum demand, "
         "average load, step-load expectations, and motor-starting needs."),
        ("Offer Type (Genset / Switchgear / Solution/Enclosure / Selective Catalytic Reduction (SCR) / Energy Storage / Other)",
         "The classification of equipment included in the Proposal, indicating whether it consists "
         "of a Genset Only or Balance of Plant (BOP) elements."),
        ("Price Validity",
         "The period during which the pricing and commercial conditions presented in this proposal "
         "remain firm before requiring review or adjustment."),
        ("Offer Purchase Order (PO) / Letter of Authorization (LOA)",
         "A formal customer document indicating financial commitment. Receipt of a PO or LOA "
         "allows scheduling and fulfillment of the proposed equipment."),
        ("Ready to Ship (RTS)",
         "The manufacturing milestone indicating that equipment is complete and available for "
         "shipment from the facility."),
        ("Scope of Supply",
         "The complete list of equipment, materials, and services included in this proposal. Items "
         "not explicitly listed are excluded."),
        ("Service Agreement (SA)",
         "Included agreement providing proactive maintenance, connectivity, and dealer support "
         "to maximize uptime."),
        ("Start on Line (SOL)",
         "A manufacturing milestone indicating when the unit formally enters the production line."),
        ("Taxes, Licenses, and Fees",
         "Duties, taxes, permits, and other governmental fees associated with procuring or installing "
         "the equipment, which are the responsibility of the Customer."),
        ("Warranty Terms",
         "The defined provisions outlining Caterpillar\u2019s obligations for repair or replacement of "
         "equipment within the specified warranty period."),
        ("Customer Value Agreement (CVA)",
         "A flexible support program designed to help owners proactively manage equipment health and "
         "long-term performance through structured access to genuine parts, service options, digital "
         "insights, and dealer expertise."),
        ("Extended Service Coverage (ESC)",
         "Caterpillar\u2019s optional long-term protection program that provides coverage for eligible "
         "component failures beyond the standard factory warranty period."),
    ]

    def_table = doc.add_table(rows=1, cols=2)
    def_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_d = def_table.rows[0]
    hdr_d.cells[0].text = "Term"
    hdr_d.cells[1].text = "Definition"
    _style_header_row(hdr_d)
    _set_table_borders(def_table)

    for term, defn in definitions:
        row = def_table.add_row()
        row.cells[0].text = term
        row.cells[1].text = defn
        for c in row.cells:
            _style_body_cell(c)

    doc.add_page_break()

    # --- Appendix B — Datasheets ----------------------------------------------
    _add_heading(doc, "Exhibit B \u2014 Datasheets", level=1)
    _add_paragraph(doc, (
        "The datasheets included in this appendix will be added by the BDM or dealer based on "
        "the specific equipment configuration selected for the curated proposal. These documents "
        "provide model-specific technical specifications such as performance ratings, fuel "
        "requirements, physical dimensions, and emissions information."
    ))
    doc.add_page_break()

    # --- Appendix C — Warranty Statement --------------------------------------
    _add_heading(doc, "Exhibit C \u2014 Warranty Statement", level=1)
    _add_paragraph(doc, (
        "Warranty terms and conditions to be provided by the authorized CAT dealer specific to "
        "this project and jurisdiction."
    ))
    doc.add_page_break()

    # --- Appendix D — Conceptual Layout ---------------------------------------
    _add_heading(doc, "Exhibit D \u2014 Conceptual Layout", level=1)

    footprint_data = _safe(sizing_result, "footprint", {})
    footprint_m2 = _safe_dict(footprint_data, "total_area_m2", None)
    if footprint_m2 is None:
        # Try as attribute
        try:
            footprint_m2 = sizing_result.footprint.total_area_m2
        except (AttributeError, TypeError):
            footprint_m2 = None

    if footprint_m2 is not None:
        try:
            fp_m2 = round(float(footprint_m2), 0)
            fp_ft2 = round(fp_m2 * 10.764, 0)
            fp_str = f"Estimated plant footprint: {fp_m2:.0f} m\u00b2 ({fp_ft2:.0f} ft\u00b2)"
        except (ValueError, TypeError):
            fp_str = "Estimated plant footprint: TBD m\u00b2 (TBD ft\u00b2)"
    else:
        fp_str = "Estimated plant footprint: TBD m\u00b2 (TBD ft\u00b2)"

    _add_paragraph(doc, fp_str, 11, bold=True)
    _add_paragraph(doc, (
        "Conceptual layout to be developed during the detailed engineering phase based on final "
        "equipment configuration and site constraints."
    ))
    doc.add_page_break()

    # --- Appendix E — Scope of Supply Matrix ----------------------------------
    _add_heading(doc, "Exhibit E \u2014 Scope of Supply Matrix", level=1)

    scope_table = doc.add_table(rows=1, cols=2)
    scope_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_s = scope_table.rows[0]
    hdr_s.cells[0].text = "In Scope"
    hdr_s.cells[1].text = "Out of Scope"
    _style_header_row(hdr_s)
    _set_table_borders(scope_table)

    in_scope = [
        "Generator set solution(s) per equipment summary",
        "Factory testing and inspection",
        "Standard documentation package",
    ]
    # Conditional in-scope items
    if use_bess and bess_power:
        in_scope.append("Battery Energy Storage System (BESS)")
    if proposal_info.get("offer_type_switchgear"):
        in_scope.append("Switchgear solution")
    if proposal_info.get("offer_type_solutions"):
        in_scope.append("Integrated solution")
    if proposal_info.get("offer_type_hybrid"):
        in_scope.append("Hybrid power configuration")

    out_scope = [
        "Civil and structural works",
        "Site preparation and foundation",
        "Installation and commissioning",
        "Grid interconnection",
        "Fuel supply infrastructure",
        "Permits and local authority approvals",
        "Taxes, duties, and import fees",
    ]

    max_rows = max(len(in_scope), len(out_scope))
    for i in range(max_rows):
        row = scope_table.add_row()
        row.cells[0].text = in_scope[i] if i < len(in_scope) else ""
        row.cells[1].text = out_scope[i] if i < len(out_scope) else ""
        for c in row.cells:
            _style_body_cell(c)

    doc.add_page_break()

    # --- Appendix F — Extended Service Coverage (ESC) -------------------------
    _add_heading(doc, "Exhibit F \u2014 Extended Service Coverage Overview", level=1)

    _add_heading(doc, "Purpose of Extended Service Coverage", level=2)
    _add_paragraph(doc, (
        "Extended Service Coverage (ESC) is Caterpillar\u2019s optional long-term protection program "
        "designed to provide coverage for eligible component failures beyond the standard factory "
        "warranty period. ESC helps protect the customer\u2019s investment by offering financial "
        "predictability and reducing the risk of unexpected repair costs over the life of the equipment."
    ))

    _add_heading(doc, "Key Customer Value", level=2)
    _add_bullet(doc, (
        "Cost Predictability: ESC converts unplanned repair expenses into a known, budgetable cost, "
        "enabling more accurate financial planning over the equipment lifecycle."
    ))
    _add_bullet(doc, (
        "Reduced Financial Risk: Coverage for major component failures helps mitigate the financial "
        "impact of unexpected equipment downtime and repair events."
    ))
    _add_bullet(doc, (
        "Lifecycle Support: ESC is designed to complement the standard warranty and extend protection "
        "into the period when equipment is most likely to require major component attention."
    ))

    _add_heading(doc, "Customer Experience Benefits", level=2)
    _add_bullet(doc, (
        "Global Coverage: ESC is backed by Caterpillar\u2019s worldwide dealer network, ensuring "
        "consistent service and support regardless of equipment location."
    ))
    _add_bullet(doc, (
        "Local Expertise: Authorized Cat dealers deliver ESC-related services using factory-trained "
        "technicians and genuine Cat parts, maintaining equipment performance and reliability."
    ))
    _add_bullet(doc, (
        "Ease of Doing Business: ESC agreements are structured to be straightforward, with clear terms "
        "and transparent coverage details that simplify long-term equipment management."
    ))
    _add_bullet(doc, (
        "Peace of Mind: Knowing that critical components are covered allows the customer to focus on "
        "core operations rather than equipment maintenance concerns."
    ))

    _add_heading(doc, "Coordination Through Your Authorized Cat Dealer", level=2)
    _add_paragraph(doc, (
        "ESC agreements are administered through the authorized Cat dealer assigned to the project. "
        "The dealer serves as the primary point of contact for coverage details, claims processing, "
        "and coordination of any required service under the agreement."
    ))

    doc.add_page_break()

    # --- Appendix G — Customer Value Agreement (CVA) --------------------------
    _add_heading(doc, "Exhibit G \u2014 Service Agreement (SA) Overview", level=1)

    _add_paragraph(doc, (
        "All Caterpillar Inc. proposals include a Service Agreement (SA) that supports "
        "Caterpillar Inc.'s brand promise by helping customers to achieve maximum uptime, "
        "lower owning and operating costs, and deliver consistent, reliable performance. "
        "The Service Agreement provides proactive maintenance, condition monitoring "
        "connectivity, and access to dealer expertise, ensuring that the asset is "
        "supported throughout its operating life."
    ))

    _add_heading(doc, "What Is a Customer Value Agreement?", level=2)
    _add_paragraph(doc, (
        "A Customer Value Agreement (CVA) is a flexible support program designed to help equipment "
        "owners proactively manage the health and long-term performance of their Cat assets. CVAs "
        "provide structured access to genuine parts, service options, digital diagnostic tools, and "
        "dealer expertise \u2014 all tailored to the specific needs of the customer and their operation."
    ))

    _add_heading(doc, "Key Value Proposition", level=2)
    _add_paragraph(doc, (
        "CVAs are built around the principle that proactive equipment management reduces total cost "
        "of ownership and maximizes uptime. By combining scheduled maintenance, condition monitoring, "
        "and expert support, CVAs help ensure that equipment operates at peak performance throughout "
        "its service life."
    ))

    _add_heading(doc, "Digital & Diagnostic Tools", level=2)
    _add_bullet(doc, (
        "S\u2022O\u2022S Fluid Analysis: Regular sampling and laboratory analysis of engine fluids "
        "to detect early signs of wear, contamination, or degradation before they lead to failures."
    ))
    _add_bullet(doc, (
        "Cat Inspections: Structured visual and technical inspections performed by dealer technicians "
        "to assess equipment condition and identify potential issues."
    ))
    _add_bullet(doc, (
        "Remote Asset Monitoring: Connected technology solutions that provide real-time visibility "
        "into equipment health, utilization, and location \u2014 enabling data-driven maintenance decisions."
    ))

    _add_heading(doc, "Dealer Expertise & Support", level=2)
    _add_paragraph(doc, (
        "CVA customers benefit from priority access to their authorized Cat dealer\u2019s service "
        "capabilities, including factory-trained technicians, genuine Cat parts, and technical "
        "support resources. The dealer works closely with the customer to develop a maintenance "
        "strategy aligned with operational goals and equipment requirements."
    ))

    _add_heading(doc, "Flexible, Customer-Focused Structure", level=2)
    _add_paragraph(doc, (
        "CVAs are designed to be modular and scalable, allowing customers to select the level of "
        "coverage and services that best fit their needs. Agreements can be adjusted over time as "
        "operational requirements change or as additional equipment is added to the fleet."
    ))

    _add_heading(doc, "Coordination Through Your Authorized Cat Dealer", level=2)
    _add_paragraph(doc, (
        "CVAs are managed through the authorized Cat dealer assigned to the project. The dealer "
        "serves as the single point of contact for all CVA-related services, including scheduling, "
        "parts ordering, inspections, and reporting."
    ))

    doc.add_page_break()

    # --- Appendix H — Additional Technical Documents --------------------------
    _add_heading(doc, "Exhibit H \u2014 Additional Technical Documents", level=1)
    _add_paragraph(doc, (
        "The materials included in this appendix will be added by the BDM or dealer based on the "
        "specific configuration in the curated proposal. This appendix may contain supplemental "
        "technical information such as detailed fuel specifications, fuel analyses, project-specific "
        "engineering assumptions, or other customer-provided documents that support the proposed "
        "solution."
    ))

    # =========================================================================
    # Serialize
    # =========================================================================
    buf = io.BytesIO()
    doc.save(buf)
    doc_bytes = buf.getvalue()

    if output_path:
        with open(output_path, "wb") as f:
            f.write(doc_bytes)

    return doc_bytes


# ---------------------------------------------------------------------------
# P41B — Public entry point (supports both old and new call signatures)
# ---------------------------------------------------------------------------

def generate_proposal_docx(
    sizing_result,
    gen_data=None,
    project_info=None,
    selected_exhibits=None,
    sizing_pdf_bytes=None,
) -> bytes:
    """
    Generate a customer proposal Word document (P41B).

    Used by render_proposal_tab in streamlit_app.py.
    For the legacy ENABLE_PROPOSAL_GEN tab, call _generate_proposal_docx_legacy directly.

    Parameters
    ----------
    sizing_result : dict
        r.model_dump() — sizing results as a plain dict.
    gen_data : dict
        Generator specs from GENERATOR_LIBRARY (iso_rating_mw, voltage_kv, ...).
    project_info : dict
        project_name, client_name, country, state_province, etc.
    selected_exhibits : list[dict]
        Optional exhibits: [{"letter": "D", "name": "Datasheets"}, ...].
    sizing_pdf_bytes : bytes, optional
        PDF bytes to reference when Sizing Report exhibit is selected.

    Returns
    -------
    bytes
        The .docx file content.
    """
    # -------------------------------------------------------------------------
    # P41B implementation — sizing_result is a dict (r.model_dump())
    # -------------------------------------------------------------------------
    if gen_data is None:
        gen_data = {}
    if project_info is None:
        project_info = {}
    if selected_exhibits is None:
        selected_exhibits = []

    # --- Extract project info -------------------------------------------------
    project_name = project_info.get("project_name") or "Untitled Project"
    client_name = project_info.get("client_name") or ""
    country = project_info.get("country") or "To be confirmed"
    state_province = project_info.get("state_province") or "To be confirmed"
    client_str = client_name if client_name else "the Customer"

    # --- Extract sizing results -----------------------------------------------
    gen_model = sizing_result.get("selected_gen", "TBD")
    n_total = sizing_result.get("n_total", "TBD")
    n_running = sizing_result.get("n_running", "")
    n_reserve = sizing_result.get("n_reserve", "")
    unit_site_cap = sizing_result.get("unit_site_cap", "TBD")
    installed_cap = sizing_result.get("installed_cap", "TBD")
    p_total_dc = sizing_result.get("p_total_dc", "TBD")
    p_it = sizing_result.get("p_it", "TBD")
    pue = sizing_result.get("pue", "TBD")
    use_bess = sizing_result.get("use_bess", False)
    bess_power = sizing_result.get("bess_power_mw", 0)
    bess_energy = sizing_result.get("bess_energy_mwh", 0)
    system_availability = sizing_result.get("system_availability", "TBD")
    total_capex = sizing_result.get("total_capex", None)
    lcoe = sizing_result.get("lcoe", None)
    freq_hz = sizing_result.get("freq_hz", "TBD")
    fuel_mode = sizing_result.get("fuel_mode", "Natural gas (pipeline)")
    site_temp_c = sizing_result.get("site_temp_c", "TBD")
    site_alt_m = sizing_result.get("site_alt_m", "TBD")
    spinning_reserve_mw = sizing_result.get("spinning_reserve_mw", "TBD")
    n_pods = sizing_result.get("n_pods", None)
    n_per_pod = sizing_result.get("n_per_pod", None)
    dc_type = sizing_result.get("dc_type", "Data Center") or "Data Center"

    # --- Generator specs from gen_data ----------------------------------------
    iso_rating = gen_data.get("iso_rating_mw", "TBD")
    voltage_kv = gen_data.get("voltage_kv", "TBD")

    # --- Format helpers -------------------------------------------------------
    def _fmt(val, fmt="{:.1f}", fallback="TBD"):
        try:
            return fmt.format(float(val))
        except (ValueError, TypeError):
            return fallback

    total_mw_str = _fmt(p_total_dc)
    derated_mw_str = _fmt(unit_site_cap, "{:.2f}")
    p_it_str = _fmt(p_it)

    if use_bess and bess_power:
        bess_str = (
            f"Yes \u2014 {_fmt(bess_power)} MW / {_fmt(bess_energy)} MWh"
        )
    else:
        bess_str = "Not included"

    try:
        avail_str = f"{float(system_availability) * 100:.2f}%"
    except (ValueError, TypeError):
        avail_str = str(system_availability)

    config_label = f"N+{n_reserve}" if n_reserve != "" else ""

    if total_capex is not None:
        try:
            formatted_capex = f"${float(total_capex):,.0f}"
        except (ValueError, TypeError):
            formatted_capex = "To be confirmed"
    else:
        formatted_capex = "To be confirmed"

    # =========================================================================
    # Create document
    # =========================================================================
    doc = Document()
    section = doc.sections[0]
    section.page_width = _PAGE_WIDTH
    section.page_height = _PAGE_HEIGHT
    section.top_margin = _MARGIN
    section.bottom_margin = _MARGIN
    section.left_margin = _MARGIN
    section.right_margin = _MARGIN

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    p_cat = doc.add_paragraph()
    p_cat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo_cover = p_cat.add_run()
    if os.path.exists(_LOGO_PATH):
        run_logo_cover.add_picture(_LOGO_PATH, width=Inches(3.5))
    else:
        run_logo_cover.text = "CATERPILLAR"
        _set_run_font(run_logo_cover, 16, bold=True)

    _add_paragraph(doc, "LEPS Global", 12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "Customer Proposal For:", 12, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, project_name, 14, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, str(gen_model), 12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, date.today().strftime("%B %d, %Y"), 11,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "Rev. 0", 11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # --- Running header from page 2 onward ------------------------------------
    new_section = doc.add_section()
    new_section.page_width = _PAGE_WIDTH
    new_section.page_height = _PAGE_HEIGHT
    new_section.top_margin = _MARGIN
    new_section.bottom_margin = _MARGIN
    new_section.left_margin = _MARGIN
    new_section.right_margin = _MARGIN
    _add_page_header(new_section, project_name)

    # =========================================================================
    # SECTION 1 — EXECUTIVE SUMMARY
    # =========================================================================
    _add_heading(doc, "Executive Summary", level=1)

    _add_paragraph(doc, (
        f"Caterpillar Inc. is pleased to present this proposal for the {project_name}. "
        "The solution offered in this document is designed to meet the project\u2019s operational, "
        "performance, and reliability requirements while providing a flexible and scalable "
        "foundation for long-term use. The proposed equipment and supporting solutions are based "
        "on proven designs widely deployed across diverse Applications and operating environments."
    ))
    _add_paragraph(doc, (
        "This document outlines a preliminary, non-binding proposal based on the information "
        "provided to date. Final pricing and commercial terms will be confirmed and executed "
        "through the authorized Caterpillar dealer."
    ))
    _add_paragraph(doc, (
        "Our organization has extensive experience supplying electric power solutions for "
        "industrial, commercial, and critical-infrastructure projects around the world. This "
        "background ensures that the solution outlined in this proposal reflects globally "
        "recognized best practices in safety, performance, and integration with customer facilities."
    ))
    _add_paragraph(doc, (
        f"Commercial information, including pricing, delivery expectations, validity, and warranty "
        f"provisions, is included within this document. The proposal is intended to provide "
        f"{client_str} with a clear understanding of the recommended solution and the associated "
        f"commercial framework."
    ))
    _add_paragraph(doc, (
        "We appreciate the opportunity to be considered for this project and trust that the "
        "information presented will support your evaluation of our proposed electric power solution."
    ), space_before=12)

    doc.add_page_break()

    # =========================================================================
    # SECTION 2 — SOLUTION OVERVIEW
    # =========================================================================
    _add_heading(doc, "Proposed Customer Solution", level=1)

    _add_paragraph(doc, (
        f"Caterpillar Inc. is pleased to provide this proposed power generation solution for "
        f"{client_str}, tailored to support the operational requirements of the {project_name}."
    ))
    _add_paragraph(doc, (
        f"This proposal aims to satisfy the request for a {total_mw_str} MW, "
        f"{dc_type} generation site utilizing the {gen_model} Caterpillar Genset."
    ))

    _add_bullet(doc, f"{n_total} x {gen_model} Caterpillar generator set solutions")
    _add_bullet(doc, f"Derated output per unit: {derated_mw_str} MW at site conditions")
    _add_bullet(doc, f"BESS: {bess_str}")
    _add_bullet(doc, f"System availability: {avail_str} ({config_label} configuration)")
    _add_bullet(doc, "CVA and ESC options available \u2014 refer to Exhibit for overview")

    doc.add_page_break()

    # =========================================================================
    # SECTION 3 — TECHNICAL OFFER
    # =========================================================================
    _add_heading(doc, "Technical Offer", level=1)

    # --- 3.1 Equipment Summary ------------------------------------------------
    _add_heading(doc, "3.1 Equipment Summary", level=2)
    _add_paragraph(doc, (
        "The Equipment Summary below outlines the equipment and selected features "
        "included in the proposed power generation solution."
    ))

    eq_table = doc.add_table(rows=1, cols=4)
    eq_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_eq = eq_table.rows[0]
    for i, label in enumerate(["Qty.", "Characteristic Name", "Feature Code", "Feature Description"]):
        hdr_eq.cells[i].text = label
    _style_header_row(hdr_eq)
    _set_table_borders(eq_table)

    try:
        feat_desc = f"{iso_rating} MW ISO, {freq_hz} Hz, {voltage_kv} kV"
    except Exception:
        feat_desc = "TBD"
    row1 = eq_table.add_row()
    row1.cells[0].text = str(n_total)
    row1.cells[1].text = "Generator Set"
    row1.cells[2].text = str(gen_model)
    row1.cells[3].text = feat_desc
    for c in row1.cells:
        _style_body_cell(c)

    if use_bess and bess_power:
        row_bess = eq_table.add_row()
        row_bess.cells[0].text = "1"
        row_bess.cells[1].text = "Battery Energy Storage System"
        row_bess.cells[2].text = "BESS"
        row_bess.cells[3].text = f"{_fmt(bess_power)} MW / {_fmt(bess_energy)} MWh"
        for c in row_bess.cells:
            _style_body_cell(c)

    for _ in range(3):
        empty_row = eq_table.add_row()
        for c in empty_row.cells:
            c.text = ""
            _style_body_cell(c)

    # --- 3.2 Project Requirements ---------------------------------------------
    _add_heading(doc, "3.2 Project Requirements", level=2)
    _add_paragraph(doc, "The following project parameters were considered in defining the technical offering:")
    _add_bullet(doc, f"Project location / site: {country}, {state_province}")
    _add_bullet(doc, "Intended application: \u2611 Prime  \u2610 Standby  \u2610 Other")
    _add_bullet(doc, (
        f"Operating environment: Ambient temperature {site_temp_c}\u00b0C, "
        f"Altitude {site_alt_m} m ASL"
    ))
    _add_bullet(doc, (
        f"Load requirements: IT Load {p_it_str} MW, PUE {pue}, "
        f"Total Facility {total_mw_str} MW"
    ))

    # --- 3.3 Fuel Requirements ------------------------------------------------
    _add_heading(doc, "3.3 Fuel Requirements", level=2)
    _add_paragraph(doc, f"The proposed generator set is designed to operate using {fuel_mode}.")
    _add_bullet(doc, f"Fuel type: {fuel_mode}")
    _add_bullet(doc, "Estimated consumption at full load: To be confirmed per fuel analysis")
    _add_bullet(doc, "Minimum inlet pressure requirements per generator specifications")

    doc.add_page_break()

    # =========================================================================
    # SECTION 4 — PRICING
    # =========================================================================
    _add_heading(doc, "Pricing", level=1)

    # --- 4.1 Base Price -------------------------------------------------------
    _add_heading(doc, "4.1 Base Price", level=2)
    _add_paragraph(doc, f"Base Price: {formatted_capex}", 11, bold=True)
    _add_paragraph(doc, "Reflects the fully integrated system scope outlined in this proposal.")
    _add_paragraph(doc, (
        "This is a budgetary estimate. Final pricing to be confirmed through "
        "authorized CAT dealer."
    ))

    # --- 4.2 Additional Options -----------------------------------------------
    _add_heading(doc, "4.2 Additional Options", level=2)
    _add_paragraph(doc, (
        "The following optional features are available upon request and are "
        "not included in the Base Price."
    ))
    opt_table = doc.add_table(rows=1, cols=3)
    opt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_opt = opt_table.rows[0]
    for i, label in enumerate(["Option", "Description", "Price (USD)"]):
        hdr_opt.cells[i].text = label
    _style_header_row(hdr_opt)
    _set_table_borders(opt_table)
    for _ in range(3):
        empty = opt_table.add_row()
        for c in empty.cells:
            c.text = ""
            _style_body_cell(c)

    # --- 4.3 Price Validity ---------------------------------------------------
    _add_heading(doc, "4.3 Price Validity", level=2)
    _add_paragraph(doc, "This proposal is valid for 90 days from date of issue.")

    # --- 4.4 Payment Terms ----------------------------------------------------
    _add_heading(doc, "4.4 Payment Terms", level=2)
    _add_paragraph(doc, "The proposed payment terms are as follows:")
    _add_bullet(doc, "30% downpayment")
    _add_bullet(doc, "30% due 30 days after purchase order")
    _add_bullet(doc, "30% due at SOL (Start on Line)")
    _add_bullet(doc, "10% due at written notice of RTS (Ready to Ship)")

    # --- 4.5 Pricing Parameters -----------------------------------------------
    _add_heading(doc, "4.5 Pricing Parameters", level=2)
    _add_paragraph(doc, (
        "a) The indicated price is based on the quoted delivery. Pricing may change "
        "if the customer requests a delay in delivery."
    ))
    _add_paragraph(doc, (
        "b) We reserve the right to adjust our prices in case some parts of our offer "
        "are inapplicable and will not be ordered."
    ))

    # --- 4.6 Taxing, Licensing, and Fees --------------------------------------
    _add_heading(doc, "4.6 Taxing, Licensing, and Fees", level=2)
    _add_paragraph(doc, (
        "a) All equipment duties, taxes of any kind, and license fees \u2014 including VAT "
        "and/or sales tax imposed by national or local government \u2014 are not included in "
        "this quotation and will be the responsibility of the Customer where applicable."
    ))
    _add_paragraph(doc, "b) Delivery per Incoterms \u2014 to be confirmed with authorized dealer.")
    _add_paragraph(doc, (
        "c) The Customer shall be responsible to secure all relevant construction permits "
        "from the national and local government agencies, and all actual fees shall be "
        "paid by the Customer."
    ))

    # --- 4.7 Delivery ---------------------------------------------------------
    _add_heading(doc, "4.7 Delivery", level=2)
    _add_paragraph(doc, (
        "a) With a Purchase Order, or document evidencing financial commitment (LOA) within "
        "the validity period stated above, the units will be available for delivery \u2014 "
        "estimated timeline to be confirmed with authorized dealer."
    ))
    _add_paragraph(doc, "b) Delivery destination: To be confirmed.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 5 — CLARIFICATIONS, DEVIATIONS & EXCEPTIONS
    # =========================================================================
    _add_heading(doc, "Clarifications, Deviations & Exceptions", level=1)

    clar_table = doc.add_table(rows=1, cols=1)
    clar_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    clar_table.rows[0].cells[0].text = "Project Clarifications, Deviations & Exceptions"
    _style_header_row(clar_table.rows[0])
    _set_table_borders(clar_table)

    boilerplate_row = clar_table.add_row()
    boilerplate_row.cells[0].text = (
        "Equipment supplied will be limited to that described in this proposal. If products "
        "and/or features, other than the equipment described in this proposal is required, "
        "please consult the factory for a re-quote as pricing may be affected. We take "
        "complete exception to any specifications that were not provided or reviewed."
    )
    _style_body_cell(boilerplate_row.cells[0])

    doc.add_page_break()

    # =========================================================================
    # SECTION 6 — COMMERCIAL ASSUMPTIONS
    # =========================================================================
    _add_heading(doc, "Commercial Assumptions", level=1)

    assumptions = [
        "All prices exclude all duties, taxes, trade surcharges, and freight.",
        (
            "Receipt of goods shall be confirmed by the customer. Notification and evidence "
            "of any material damaged during shipment and/or any concerns with the products "
            "shall be provided as soon as possible but not later than 10 business days "
            "following shipment loading or 5 days after receipt, whichever comes first."
        ),
        (
            "Unless specifically indicated, all job site installation, commissioning and "
            "testing for the project is NOT included."
        ),
        (
            "Unless specifically indicated, storage of offered materials at the manufacturing "
            "location is NOT available. Completed units must be picked up from the manufacturing "
            "facility within 5 business days of RTS. Units that are not picked up within 5 "
            "business days will be moved to an off-site storage facility at customer\u2019s expense."
        ),
        (
            "Any modification to the defined scope of supply requires a mutually agreed and "
            "signed change order."
        ),
        (
            "Caterpillar product warranty is governed by the applicable Caterpillar Inc. "
            "Warranty Guide(s)."
        ),
        (
            "Orders may be cancelled only by written notice. Payments made under Payment Terms "
            "are non-refundable and will not be returned in event of cancellation."
        ),
    ]

    assump_table = doc.add_table(rows=1, cols=2)
    assump_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_a = assump_table.rows[0]
    hdr_a.cells[0].text = "No."
    hdr_a.cells[1].text = "Description"
    _style_header_row(hdr_a)
    _set_table_borders(assump_table)

    for idx, assumption_text in enumerate(assumptions, 1):
        row_a = assump_table.add_row()
        row_a.cells[0].text = str(idx)
        row_a.cells[1].text = assumption_text
        for c in row_a.cells:
            _style_body_cell(c)

    doc.add_page_break()

    # =========================================================================
    # SECTION 7 — APPENDICES (MANDATORY A, B, C + DYNAMIC D, E, ...)
    # =========================================================================

    # --- Exhibit A — Definitions (MANDATORY) ----------------------------------
    _add_heading(doc, "Exhibit A \u2014 Definitions", level=1)

    definitions = [
        ("Application (Standby / Prime / Other)",
         "The operational mode of the generator set. Standby is for emergency use only during "
         "utility outages, while Prime is intended for continuous or frequent operation as the "
         "main power source."),
        ("Balance of Plant (BOP)",
         "All supporting equipment required for a complete power solution beyond the generator set "
         "itself. In this Proposal, BOP includes switchgear, enclosures/solutions, SCR solutions, "
         "inverters, energy storage, and other ancillary components."),
        ("Base Price",
         "The total price for the defined Scope of Supply included in the Proposal, excluding "
         "optional items, taxes, duties, freight, and licensing fees."),
        ("Bill of Material (BOM)",
         "A summarized list of all equipment, components, features, and selected options included "
         "in the proposed power generation solution."),
        ("Clarification",
         "An explanation of assumptions or interpretations where project information is incomplete "
         "or ambiguous, provided to ensure shared understanding of the proposal."),
        ("Delivery",
         "A delivery condition in which Caterpillar makes the equipment available at the factory, "
         "and the Customer assumes responsibility for transportation, duties, permits, and logistics "
         "from that point forward."),
        ("Deviation",
         "A departure from the customer\u2019s specification when the proposed solution cannot fully "
         "conform to the requirements provided."),
        ("Exception",
         "A requested requirement or obligation that Caterpillar cannot comply with or accept as "
         "part of this proposal."),
        ("Feature Code",
         "The identifier used to specify a configurable feature, option, or characteristic of the "
         "generator set or BOP equipment."),
        ("Fuel Requirements",
         "The fuel type and characteristics required for operation of the generator set, such as "
         "natural gas, biogas, diesel, or propane."),
        ("Load Requirements",
         "The expected electrical demand characteristics of the project, including maximum demand, "
         "average load, step-load expectations, and motor-starting needs."),
        ("Offer Type (Genset / Switchgear / Solution/Enclosure / SCR / Energy Storage / Other)",
         "The classification of equipment included in the Proposal, indicating whether it consists "
         "of a Genset Only or Balance of Plant (BOP) elements."),
        ("Price Validity",
         "The period during which the pricing and commercial conditions presented in this proposal "
         "remain firm before requiring review or adjustment."),
        ("Purchase Order (PO) / Letter of Authorization (LOA)",
         "A formal customer document indicating financial commitment. Receipt of a PO or LOA "
         "allows scheduling and fulfillment of the proposed equipment."),
        ("Ready to Ship (RTS)",
         "The manufacturing milestone indicating that equipment is complete and available for "
         "shipment from the facility."),
        ("Scope of Supply",
         "The complete list of equipment, materials, and services included in this proposal. Items "
         "not explicitly listed are excluded."),
        ("Service Agreement (SA)",
         "An agreement providing proactive maintenance, connectivity, and dealer support "
         "to maximize uptime."),
        ("Start on Line (SOL)",
         "A manufacturing milestone indicating when the unit formally enters the production line."),
        ("Taxes, Licenses, and Fees",
         "Duties, taxes, permits, and other governmental fees associated with procuring or installing "
         "the equipment, which are the responsibility of the Customer."),
        ("Warranty Terms",
         "The defined provisions outlining Caterpillar\u2019s obligations for repair or replacement of "
         "equipment within the specified warranty period."),
        ("Customer Value Agreement (CVA)",
         "A flexible support program designed to help owners proactively manage equipment health and "
         "long-term performance through structured access to genuine parts, service options, digital "
         "insights, and dealer expertise."),
        ("Extended Service Coverage (ESC)",
         "Caterpillar\u2019s optional long-term protection program that provides coverage for eligible "
         "component failures beyond the standard factory warranty period."),
    ]

    def_table = doc.add_table(rows=1, cols=2)
    def_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_d = def_table.rows[0]
    hdr_d.cells[0].text = "Term"
    hdr_d.cells[1].text = "Definition"
    _style_header_row(hdr_d)
    _set_table_borders(def_table)

    for term, defn in definitions:
        row_d = def_table.add_row()
        row_d.cells[0].text = term
        row_d.cells[1].text = defn
        for c in row_d.cells:
            _style_body_cell(c)

    doc.add_page_break()

    # --- Exhibit B — Extended Service Coverage (MANDATORY) --------------------
    _add_heading(doc, "Exhibit B \u2014 Extended Service Coverage Overview", level=1)

    _add_heading(doc, "Purpose of Extended Service Coverage", level=2)
    _add_paragraph(doc, (
        "Extended Service Coverage (ESC) is Caterpillar\u2019s optional long-term protection program "
        "designed to provide coverage for eligible component failures beyond the standard factory "
        "warranty period. ESC helps protect the customer\u2019s investment by offering financial "
        "predictability and reducing the risk of unexpected repair costs over the life of the equipment."
    ))

    _add_heading(doc, "Key Customer Value", level=2)
    _add_bullet(doc, (
        "Cost Predictability: ESC converts unplanned repair expenses into a known, budgetable cost, "
        "enabling more accurate financial planning over the equipment lifecycle."
    ))
    _add_bullet(doc, (
        "Reduced Financial Risk: Coverage for major component failures helps mitigate the financial "
        "impact of unexpected equipment downtime and repair events."
    ))
    _add_bullet(doc, (
        "Lifecycle Support: ESC extends protection into the period when equipment is most likely "
        "to require major component attention."
    ))

    _add_heading(doc, "Customer Experience Benefits", level=2)
    _add_bullet(doc, (
        "Global Coverage: ESC is backed by Caterpillar\u2019s worldwide dealer network, ensuring "
        "consistent service and support regardless of equipment location."
    ))
    _add_bullet(doc, (
        "Local Expertise: Authorized Cat dealers deliver ESC-related services using factory-trained "
        "technicians and genuine Cat parts."
    ))
    _add_bullet(doc, (
        "Peace of Mind: Knowing that critical components are covered allows the customer to focus "
        "on core operations rather than equipment maintenance concerns."
    ))

    _add_heading(doc, "Coordination Through Your Authorized Cat Dealer", level=2)
    _add_paragraph(doc, (
        "ESC agreements are administered through the authorized Cat dealer assigned to the project. "
        "The dealer serves as the primary point of contact for coverage details, claims processing, "
        "and coordination of any required service under the agreement."
    ))

    doc.add_page_break()

    # --- Exhibit C — Customer Value Agreement (MANDATORY) --------------------
    _add_heading(doc, "Exhibit C \u2014 Service Agreement (SA) Overview", level=1)

    _add_paragraph(doc, (
        "All Caterpillar Inc. proposals include a Service Agreement (SA) that supports "
        "Caterpillar Inc.\u2019s brand promise by helping customers to achieve maximum uptime, "
        "lower owning and operating costs, and deliver consistent, reliable performance. "
        "The Service Agreement provides proactive maintenance, condition monitoring "
        "connectivity, and access to dealer expertise, ensuring that the asset is "
        "supported throughout its operating life."
    ))

    _add_heading(doc, "What Is a Customer Value Agreement?", level=2)
    _add_paragraph(doc, (
        "A Customer Value Agreement (CVA) is a flexible support program designed to help equipment "
        "owners proactively manage the health and long-term performance of their Cat assets. CVAs "
        "provide structured access to genuine parts, service options, digital diagnostic tools, and "
        "dealer expertise \u2014 all tailored to the specific needs of the customer and their operation."
    ))

    _add_heading(doc, "Key Value Proposition", level=2)
    _add_paragraph(doc, (
        "CVAs are built around the principle that proactive equipment management reduces total cost "
        "of ownership and maximizes uptime. By combining scheduled maintenance, condition monitoring, "
        "and expert support, CVAs help ensure that equipment operates at peak performance throughout "
        "its service life."
    ))

    _add_heading(doc, "Digital & Diagnostic Tools", level=2)
    _add_bullet(doc, (
        "S\u2022O\u2022S Fluid Analysis: Regular sampling and laboratory analysis of engine fluids "
        "to detect early signs of wear, contamination, or degradation before they lead to failures."
    ))
    _add_bullet(doc, (
        "Cat Inspections: Structured visual and technical inspections performed by dealer technicians "
        "to assess equipment condition and identify potential issues."
    ))
    _add_bullet(doc, (
        "Remote Asset Monitoring: Connected technology solutions that provide real-time visibility "
        "into equipment health, utilization, and location \u2014 enabling data-driven maintenance decisions."
    ))

    _add_heading(doc, "Coordination Through Your Authorized Cat Dealer", level=2)
    _add_paragraph(doc, (
        "CVAs are managed through the authorized Cat dealer assigned to the project. The dealer "
        "serves as the single point of contact for all CVA-related services, including scheduling, "
        "parts ordering, inspections, and reporting."
    ))

    doc.add_page_break()

    # =========================================================================
    # DYNAMIC EXHIBITS (D, E, F, ... based on selected_exhibits)
    # =========================================================================
    for exhibit in selected_exhibits:
        letter = exhibit.get("letter", "?")
        name = exhibit.get("name", "")

        _add_heading(doc, f"Exhibit {letter} \u2014 {name}", level=1)

        if name == "Datasheets":
            _add_paragraph(doc, (
                f"Equipment datasheets to be added by BDM/dealer for {gen_model}. "
                "These documents provide model-specific technical specifications including "
                "performance ratings, fuel requirements, physical dimensions, and emissions information."
            ))

        elif name == "Warranty Statement":
            _add_paragraph(doc, (
                "Warranty terms to be provided by the authorized Cat dealer for this project "
                "and jurisdiction."
            ))

        elif name == "Conceptual Layout":
            _add_paragraph(doc, (
                "Conceptual site layout to be provided during the detailed engineering phase "
                "based on final equipment configuration and site constraints."
            ))
            footprint = sizing_result.get("footprint")
            if isinstance(footprint, dict):
                fp_m2 = footprint.get("total_area_m2")
                if fp_m2 is not None:
                    try:
                        fp_ft2 = float(fp_m2) * 10.764
                        _add_paragraph(
                            doc,
                            f"Estimated plant footprint: {float(fp_m2):.0f} m\u00b2 "
                            f"({fp_ft2:.0f} ft\u00b2)",
                            bold=True,
                        )
                    except (ValueError, TypeError):
                        pass

        elif name == "Scope of Supply Matrix":
            _add_paragraph(doc, "Scope of supply matrix to be provided with final engineering package.")
            scope_tbl = doc.add_table(rows=1, cols=2)
            scope_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_sc = scope_tbl.rows[0]
            hdr_sc.cells[0].text = "In Scope"
            hdr_sc.cells[1].text = "Out of Scope"
            _style_header_row(hdr_sc)
            _set_table_borders(scope_tbl)
            in_scope_items = [
                f"{n_total} x {gen_model} Generator Set Solution(s)",
                "Factory testing and inspection",
                "Standard documentation package",
            ]
            if use_bess and bess_power:
                in_scope_items.append("Battery Energy Storage System (BESS)")
            out_scope_items = [
                "Civil and structural works",
                "Site preparation and foundation",
                "Installation and commissioning",
                "Grid interconnection",
                "Fuel supply infrastructure",
                "Permits and local authority approvals",
                "Taxes, duties, and import fees",
            ]
            max_rows = max(len(in_scope_items), len(out_scope_items))
            for i in range(max_rows):
                row_sc = scope_tbl.add_row()
                row_sc.cells[0].text = in_scope_items[i] if i < len(in_scope_items) else ""
                row_sc.cells[1].text = out_scope_items[i] if i < len(out_scope_items) else ""
                for c in row_sc.cells:
                    _style_body_cell(c)

        elif name == "Sizing Report":
            _add_paragraph(doc, (
                "This exhibit presents a summary of the power system sizing analysis performed "
                "for this project. Detailed calculations are available upon request."
            ))
            _add_heading(doc, "System Configuration Summary", level=2)
            sr_table = doc.add_table(rows=1, cols=2)
            sr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_sr = sr_table.rows[0]
            hdr_sr.cells[0].text = "Parameter"
            hdr_sr.cells[1].text = "Value"
            _style_header_row(hdr_sr)
            _set_table_borders(sr_table)

            sizing_summary_rows = [
                ("Generator Model", str(gen_model)),
                ("Fleet Size (Total / Running / Reserve)", f"{n_total} / {n_running} / {n_reserve}"),
                ("Installed Capacity", f"{_fmt(installed_cap)} MW"),
                ("Derated Output per Unit", f"{derated_mw_str} MW"),
                ("IT Load", f"{p_it_str} MW"),
                ("Total Facility Load (incl. PUE)", f"{total_mw_str} MW"),
                ("PUE", str(pue)),
                ("Site Temperature", f"{site_temp_c} \u00b0C"),
                ("Site Altitude", f"{site_alt_m} m ASL"),
                ("Fuel Type", str(fuel_mode)),
                ("Spinning Reserve", f"{_fmt(spinning_reserve_mw)} MW"),
                ("System Availability", avail_str),
                ("BESS", bess_str),
            ]
            if total_capex is not None:
                sizing_summary_rows.append(("Total CAPEX (Budgetary)", formatted_capex))
            if lcoe is not None:
                try:
                    sizing_summary_rows.append(("LCOE", f"${float(lcoe):.2f} / MWh"))
                except (ValueError, TypeError):
                    pass
            if n_pods is not None and n_per_pod is not None:
                sizing_summary_rows.append(
                    ("Pod Architecture", f"{n_pods} pods \u00d7 {n_per_pod} units/pod")
                )

            for param, value in sizing_summary_rows:
                row_sr = sr_table.add_row()
                row_sr.cells[0].text = param
                row_sr.cells[1].text = value
                for c in row_sr.cells:
                    _style_body_cell(c)

            if sizing_pdf_bytes:
                _add_paragraph(doc, (
                    "Note: Full sizing report PDF is attached separately. "
                    "See attached document for complete calculation details."
                ), space_before=12)

        elif name == "Additional Technical Documents":
            _add_paragraph(doc, (
                "Additional technical documents to be provided by the BDM or dealer based on "
                "the specific configuration in the curated proposal. This exhibit may contain "
                "supplemental technical information such as detailed fuel specifications, "
                "fuel analyses, project-specific engineering assumptions, or other "
                "customer-provided documents that support the proposed solution."
            ))

        else:
            _add_paragraph(doc, f"Content for {name} to be provided.")

        doc.add_page_break()

    # =========================================================================
    # Serialize
    # =========================================================================
    p41b_buf = io.BytesIO()
    doc.save(p41b_buf)
    p41b_bytes = p41b_buf.getvalue()

    if output_path:
        with open(output_path, "wb") as f:
            f.write(p41b_bytes)

    return p41b_bytes
